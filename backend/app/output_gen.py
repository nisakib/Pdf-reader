import os
import logging
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

logger = logging.getLogger(__name__)


def generate_docx(pages: list[dict], output_path: str) -> None:
    """Generate a DOCX file with translated English text per page."""
    doc = Document()
    doc.add_heading("Japanese → English Translation", level=1)

    for page in pages:
        doc.add_heading(f"Page {page['page_number']}", level=2)
        doc.add_paragraph(page["english_text"] or "(no content)")
        doc.add_paragraph()  # blank line

    doc.save(output_path)
    logger.info(f"DOCX saved to {output_path}")


def generate_pdf(pages: list[dict], output_path: str) -> None:
    """Generate a PDF file with translated English text per page."""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontSize=16,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "Heading2",
        parent=styles["Heading2"],
        fontSize=13,
        spaceAfter=8,
        spaceBefore=16,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=6,
        leading=14,
    )

    story = [
        Paragraph("Japanese → English Translation", title_style),
        Spacer(1, 6 * mm),
    ]

    for page in pages:
        story.append(Paragraph(f"Page {page['page_number']}", heading_style))
        text = (page["english_text"] or "(no content)").replace("\n", "<br/>")
        story.append(Paragraph(text, body_style))
        story.append(Spacer(1, 4 * mm))

    doc.build(story)
    logger.info(f"PDF saved to {output_path}")
